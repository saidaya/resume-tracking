# Importing Flask Lib
from flask import Blueprint, jsonify, request
from flask import Flask, request, send_file
from flask_cors import CORS, cross_origin

# Import service and model

from resume_jd.resume_jd_service import extract_skills_from_resume
from resume_jd.resume_jd_service  import extract_skills_jd
from resume_jd.resume_jd_service import scorer
from resume_jd.resume_jd_service import matchingScore
from resume_jd.resume_jd_model import ResponseData

#Importing Lib for JSON,DOCX,PDF
import json
import PyPDF2
from io import BytesIO
import docx
import os


# PATH for API route and skills JSON
resume_jd_route_path = 'resume_jd/v1'
resume_jd_route = Blueprint(resume_jd_route_path, __name__)



#API to get the skills from JD - new
@resume_jd_route.route("/jd_skills", methods=['POST'])
def skills_jd():
    data = request.get_json()
    job_description = data.get('jd')
    company_name = data.get('company_name')
    skills = extract_skills_jd(job_description, company_name)
    if "status" in skills:
        # If extract_skills_jd_1 function returns an error
        response_data = {
            'company_name': company_name,
            'job_description': job_description,
            'message': skills["message"]
        }
        return jsonify(response_data), 500  # Internal Server Error status code

    # If extract_skills_jd_1 function returns skills data
    response_data = {
        'company_name': company_name,
        'job_description': job_description,
        'skills': skills['skills']
    }
    print(response_data)
    return jsonify(response_data)

#API to get the skills from Resume -new
@resume_jd_route.route('/resume_skills', methods=['POST'])
def skills_resume():
    # Check if the file was uploaded
    if 'file' not in request.files:
        response_data = {
            'message': 'No file part in the request'
        }
        return jsonify(response_data), 400

    uploaded_file = request.files['file']
    file_name = uploaded_file.filename

    # Check if the filename is valid
    if not file_name or '.' not in file_name or file_name.rsplit('.', 1)[1].lower() not in ['pdf', 'docx', 'txt']:
        response_data = {
            'message': 'Invalid file format'
        }
        return jsonify(response_data), 400

    # Read the uploaded file based on its type
    if file_name.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
        resume_text = ''
        for page in pdf_reader.pages:
            resume_text += page.extract_text()
    elif file_name.endswith('.docx'):
        doc = docx.Document(BytesIO(uploaded_file.read()))
        resume_text = ''
        for paragraph in doc.paragraphs:
            resume_text += paragraph.text
    elif file_name.endswith('.txt'):
        resume_text = uploaded_file.read().decode('utf-8')
    else:
        response_data = {
            'message': 'Invalid file format'
        }
        return jsonify(response_data), 400

    # Extract skills from the resume text
    skills_response = extract_skills_from_resume(resume_text)

    # Check if there was an error extracting skills
    if "status" in skills_response and skills_response["status"] != 200:
        return jsonify(skills_response), skills_response["status"]

    # Prepare the response
    response_data = {
        'message': f'{file_name} is uploaded successfully',
        'resume_text': resume_text,
        'skills': skills_response["skills"]
    }
    return jsonify(response_data), 200

#API to get cosine Simmilarity: - new
@resume_jd_route.route('/skills/score', methods=['GET'])
def get_score_new():
    try:
        cosine_score = scorer()['cosine_similarity']
        return jsonify({
            'message': 'Success',
            'score': cosine_score,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500




@resume_jd_route.route('/skills/score/files', methods=['POST'])
def get_score():
    try:
        job_description = request.form['jd']
        resume_file = request.files['resume']


        # Check if resume file exists in the request
        if not resume_file:
            response_data = {'message': 'No resume file in the request'}
            return jsonify(response_data), 400

        # Check if the file format is valid
        file_name = resume_file.filename
        if not file_name or '.' not in file_name or file_name.rsplit('.', 1)[1].lower() not in ['pdf', 'docx', 'txt']:
            response_data = {'message': 'Invalid file format'}
            return jsonify(response_data), 400

        # Read the uploaded file based on its type
        if file_name.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(BytesIO(resume_file.read()))
            resume_text = ''
            for page in pdf_reader.pages:
                resume_text += page.extract_text()
        elif file_name.endswith('.docx'):
            doc = docx.Document(BytesIO(resume_file.read()))
            resume_text = ''
            for paragraph in doc.paragraphs:
                resume_text += paragraph.text
        elif file_name.endswith('.txt'):
            resume_text = resume_file.read().decode('utf-8')
        else:
            response_data = {'message': 'Invalid file format'}
            return jsonify(response_data), 400




        # Calculate cosine similarity score between job description skills and resume skills
        skills_score = matchingScore(job_description,resume_text)['cosine_similarity']

        # Prepare the response
        response_data = {
            'message': 'Success',
            'skills_score': skills_score
        }
        return jsonify(response_data), 200

    except Exception as e:
        response_data = {'error': str(e)}
        return jsonify(response_data), 500

